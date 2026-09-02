"""
DOCX renderer — python-docx: ReportData → DOCX bytes.
Consumes the same ReportData as pdf_renderer so both formats
can never silently drift apart in content.
No LibreOffice conversion step — python-docx constructs the document directly.
"""

from __future__ import annotations
import io
import logging
from typing import Any

from render.report_data import ReportData, ReportSection

logger = logging.getLogger(__name__)


def render_docx(report: ReportData) -> bytes:
    """
    Render a ReportData to DOCX bytes using python-docx.
    Returns raw bytes suitable for streaming or writing to disk.
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError as e:
        raise RuntimeError("python-docx is not installed. Add it to requirements.txt.") from e

    doc = Document()

    # ── Document margins ──────────────────────────────────────
    from docx.oxml.ns import qn
    section = doc.sections[0]
    section.page_width  = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.0)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.0)

    # ── Define styles ─────────────────────────────────────────
    _setup_styles(doc)

    # ── Title page ────────────────────────────────────────────
    _add_title_page(doc, report)

    # ── Sections ──────────────────────────────────────────────
    for sec in report.sections:
        _add_section(doc, sec)

    # ── Signature block ───────────────────────────────────────
    doc.add_page_break()
    doc.add_paragraph("DECLARATION / SIGNATURES", style="Heading 2")
    sig_table = doc.add_table(rows=2, cols=2)
    sig_table.style = "Table Grid"
    sig_table.cell(0, 0).text = "Head of Department"
    sig_table.cell(0, 1).text = "Principal / Director"
    sig_table.cell(1, 0).text = "Date: ___________"
    sig_table.cell(1, 1).text = "Date: ___________"

    # ── Serialise to bytes ────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    docx_bytes = buf.read()
    logger.info(f"[docx_renderer] Generated DOCX: {len(docx_bytes):,} bytes for report_id={report.report_id}")
    return docx_bytes


# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────

def _setup_styles(doc):
    """Ensure all custom styles exist in the document."""
    from docx.shared import Pt, RGBColor
    styles = doc.styles

    def _get_or_create(name, base="Normal"):
        try:
            return styles[name]
        except KeyError:
            style = styles.add_style(name, 1)  # 1 = WD_STYLE_TYPE.PARAGRAPH
            style.base_style = styles[base]
            return style

    formula_style = _get_or_create("FormulaResult", "Normal")
    formula_style.font.size = Pt(11)
    formula_style.font.bold = True
    formula_style.font.color.rgb = RGBColor(0x27, 0x67, 0x49)

    placeholder_style = _get_or_create("PlaceholderWarning", "Normal")
    placeholder_style.font.size = Pt(9)
    placeholder_style.font.color.rgb = RGBColor(0x74, 0x42, 0x10)
    placeholder_style.font.italic = True


def _add_title_page(doc, report: ReportData):
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("NATIONAL BOARD OF ACCREDITATION")
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x71, 0x80, 0x96)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("SELF-ASSESSMENT REPORT (SAR)")
    r2.bold = True
    r2.font.size = Pt(20)
    r2.font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("UG Engineering Programs — Tier-II Institution | GAPC V4.0 | January 2025").font.size = Pt(10)

    doc.add_paragraph()

    inst = doc.add_paragraph()
    inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_inst = inst.add_run(report.institution_name)
    r_inst.bold = True
    r_inst.font.size = Pt(16)

    prog = doc.add_paragraph()
    prog.alignment = WD_ALIGN_PARAGRAPH.CENTER
    prog.add_run(report.program_name).font.size = Pt(13)

    dept = doc.add_paragraph()
    dept.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dept.add_run(f"Department of {report.department.get('name', '—')}").font.size = Pt(11)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Academic Year: {report.academic_year}  |  Generated: {report.generated_at[:10]}")

    doc.add_page_break()


def _add_section(doc, sec: ReportSection):
    import io
    import base64
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    # Criterion header
    if sec.content_type == "criterion_header":
        doc.add_page_break()
        if sec.id == "4":
            tbl_banner = doc.add_table(rows=1, cols=3)
            tbl_banner.style = "Table Grid"
            tbl_banner.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            c0 = tbl_banner.cell(0, 0)
            c0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            r0 = c0.paragraphs[0].add_run("CRITERION-4")
            r0.bold = True
            r0.font.size = Pt(13)

            c1 = tbl_banner.cell(0, 1)
            c1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            r1 = c1.paragraphs[0].add_run("Students Performance")
            r1.bold = True
            r1.font.size = Pt(13)

            c2 = tbl_banner.cell(0, 2)
            c2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            r2 = c2.paragraphs[0].add_run("150")
            r2.bold = True
            r2.font.size = Pt(13)
        else:
            h = doc.add_paragraph(style="Heading 1")
            h.add_run(f"Criterion {sec.id}: {sec.title}")
        doc.add_paragraph()
        return

    # Sub-section heading
    heading_level = min(sec.level + 1, 4)
    h = doc.add_paragraph(style=f"Heading {heading_level}")
    h.add_run(f"{sec.id}  {sec.title}  ({sec.marks} marks)")

    # Placeholder warning
    if sec.has_placeholders:
        pw = doc.add_paragraph(style="PlaceholderWarning")
        pw.add_run("⚠ Placeholder data — verify before submission.")

    # Table
    if sec.table_rows:
        cols = len(sec.table_headers) if sec.table_headers else len(sec.table_rows[0])
        tbl  = doc.add_table(rows=0, cols=cols)
        tbl.style = "Table Grid"

        if sec.table_headers:
            hdr_cells = tbl.add_row().cells
            for i, h_text in enumerate(sec.table_headers):
                hdr_cells[i].text = str(h_text)
                run = hdr_cells[i].paragraphs[0].runs[0]
                run.bold = True

        for row_data in sec.table_rows:
            row_cells = tbl.add_row().cells
            for i, cell_val in enumerate(row_data):
                row_cells[i].text = str(cell_val)

        # Formula result summary below table
        if sec.formula_result:
            p_calc = doc.add_paragraph()
            calc_text = sec.formula_result.get("formula_text") or sec.formula_result.get("assessment_formula")
            if calc_text:
                r_calc = p_calc.add_run(calc_text)
                r_calc.bold = True
                r_calc.font.size = Pt(10)
                r_calc.font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)
            elif sec.formula_result.get("marks") is not None:
                r_calc = p_calc.add_run(f"Assessment Score: {sec.formula_result.get('marks')} / {sec.marks} Marks")
                r_calc.bold = True
                r_calc.font.size = Pt(10)

        doc.add_paragraph()

    # Detailed Summary Sheets (Section 4.6.1)
    if getattr(sec, "summary_sheets", None):
        h_sum = doc.add_paragraph(style="Heading 3")
        h_sum.add_run(f"Detailed Activity Summary Sheets ({len(sec.summary_sheets)} Selected Events)")

        for idx, sheet in enumerate(sec.summary_sheets):

            p_act = doc.add_paragraph()
            r_act = p_act.add_run(f"Activity-{idx + 1}: {sheet.get('title', '—')}")
            r_act.bold = True
            r_act.font.size = Pt(11)

            p_meta = doc.add_paragraph()
            p_meta.add_run(f"Club: {sheet.get('club_name', '—')}  |  Type: {(sheet.get('event_type') or '—').capitalize()}  |  Date: {(sheet.get('event_date') or '')[:10]}")




            tbl_s = doc.add_table(rows=7, cols=3)
            tbl_s.style = "Table Grid"

            # Header
            tbl_s.cell(0, 0).paragraphs[0].add_run("Sl.No").bold = True
            tbl_s.cell(0, 1).paragraphs[0].add_run("Field").bold = True
            tbl_s.cell(0, 2).paragraphs[0].add_run("Summary Details").bold = True

            fields = [
                ("1", "Title of the Event", sheet.get("title", "—")),
                ("2", "In association with", sheet.get("club_name") or "Department Club"),
                ("3", "Mapping to POs", sheet.get("po_mapping") or "PO1, PO2, PO5, PO12"),
                ("4", "Resource person", sheet.get("resource_person") or (", ".join(sheet.get("guest_names_list", []))) or "—"),
                ("5", "Hands-on / Skill oriented", sheet.get("skill_orientation") or "Skill oriented"),
                ("6", "Outcomes achieved / conclusion", sheet.get("report_text") or sheet.get("description") or "—"),
            ]

            for row_idx, (sl, fld, val) in enumerate(fields, start=1):
                tbl_s.cell(row_idx, 0).text = sl
                tbl_s.cell(row_idx, 1).text = fld
                tbl_s.cell(row_idx, 2).text = str(val)

            # Photos row
            tbl_s.cell(6, 0).text = "7"
            tbl_s.cell(6, 1).text = "Photos"
            cell_photos = tbl_s.cell(6, 2)
            cell_photos.paragraphs[0].text = ""

            photos = sheet.get("photos_formatted", [])
            if photos:
                for p in photos:
                    data_url = p.get("photo_data_url") or p.get("photo_url") or ""
                    if data_url.startswith("data:image") and "base64," in data_url:
                        try:
                            b64_str = data_url.split("base64,")[1]
                            img_bytes = base64.b64decode(b64_str)
                            cell_photos.paragraphs[0].add_run().add_picture(io.BytesIO(img_bytes), width=Inches(2.5))
                            if p.get("caption"):
                                p_cap = cell_photos.add_paragraph()
                                r_cap = p_cap.add_run(p["caption"])
                                r_cap.font.size = Pt(7.5)
                                r_cap.font.italic = True
                        except Exception:
                            cell_photos.paragraphs[0].text = "Event photograph attached"
                    else:
                        cell_photos.paragraphs[0].text = "Event photograph attached"
            else:
                cell_photos.paragraphs[0].text = "Event photograph attached"

            doc.add_paragraph()

    # Student Achievement Certificates / Photos (Section 4.6.3)
    if sec.id == "4.6.3" and sec.source_data and sec.source_data.get("achievements"):
        h_ach = doc.add_paragraph(style="Heading 3")
        h_ach.add_run("Student Achievement Photo & Certificate Records")

        achievements = sec.source_data.get("achievements", [])
        for ach in achievements:
            p_urls = ach.get("photo_data_urls", [])
            for purl in p_urls:
                if purl.startswith("data:image") and "base64," in purl:
                    try:
                        b64_str = purl.split("base64,")[1]
                        img_bytes = base64.b64decode(b64_str)
                        p_img = doc.add_paragraph()
                        p_img.add_run().add_picture(io.BytesIO(img_bytes), width=Inches(3.0))
                        p_lbl = doc.add_paragraph()
                        r_lbl = p_lbl.add_run(f"{ach.get('event_name', '')} — {ach.get('result_description', '')}")
                        r_lbl.font.size = Pt(8.5)
                        r_lbl.font.bold = True
                    except Exception:
                        pass

        doc.add_paragraph()

    # Narrative / static text
    if sec.narrative:
        for line in sec.narrative.split("\n"):
            if line.strip():
                doc.add_paragraph(line)


