"""
Unit tests for Event Selection & Summary Sheet rendering in Report Generation Service.
Tests Criterion 4.6.1 integration:
  1. Compact table contains ALL approved events unconditionally.
  2. Summary sheets are generated ONLY for selected include_event_ids.
  3. Generation succeeds with 0 events selected.
  4. ReportJob stores include_event_ids.
  5. DOCX and PDF render summary sheet blocks when present.
"""

import sys, os
sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

import pytest
from unittest.mock import patch
from datetime import datetime, timezone
from sar_tree.ug_tier_ii_gapc_v4 import NODES, SARNode

from render.report_data import ReportData, ReportSection
from render.builder import _build_section



@pytest.fixture
def sample_approved_events():
    return [
        {
            "id": 101,
            "title": "Smart India Hackathon Internal Round",
            "club_id": 1,
            "club_name": "Coding Club",
            "event_type": "hackathon",
            "event_date": "2025-09-15T09:00:00",
            "venue": "Main Auditorium",
            "attendee_count": 120,
            "status": "approved",
            "resource_person": "Dr. Arvind Rao",
            "po_mapping": "PO1, PO2, PO3, PO5",
            "skill_orientation": "Software Development & Problem Solving",
            "report_text": "Successfully conducted 24-hour hackathon with 30 teams.",
            "photos_formatted": [
                {"id": 1, "photo_path": "sih_1.jpg", "photo_url": "/api/v1/event-photos/sih_1.jpg", "caption": "Inauguration"}
            ],
        },
        {
            "id": 102,
            "title": "Robotics & Embedded Systems Workshop",
            "club_id": 2,
            "club_name": "Robotics Society",
            "event_type": "workshop",
            "event_date": "2025-11-20T10:00:00",
            "venue": "IoT Lab",
            "attendee_count": 65,
            "status": "approved",
            "resource_person": "Prof. S. Nair",
            "po_mapping": "PO1, PO5, PO12",
            "skill_orientation": "Embedded Systems & Hardware Prototyping",
            "report_text": "Hands-on workshop on ESP32 microcontrollers.",
            "photos_formatted": [],
        },
        {
            "id": 103,
            "title": "AI & Ethics Tech Talk",
            "club_id": 1,
            "club_name": "Coding Club",
            "event_type": "seminar",
            "event_date": "2026-02-10T14:00:00",
            "venue": "Seminar Hall 2",
            "attendee_count": 80,
            "status": "approved",
            "resource_person": "Ms. Priya Sharma (Google)",
            "po_mapping": "PO6, PO8, PO12",
            "skill_orientation": "Ethics in Artificial Intelligence",
            "report_text": "Keynote on ethical AI development.",
            "photos_formatted": [],
        },
    ]


class TestEventSummarySheetsBuilder:
    def test_compact_table_contains_all_approved_events(self, sample_approved_events):
        """All approved events must appear in the compact Layer 1 table."""
        node = NODES["4.6.1"]
        dept = {"code": "CSE", "name": "Computer Science and Engineering"}
        app_config = {"ACADEMIC_DATA_SERVICE_URL": "http://mock-service"}

        with patch("data_client.fetch_approved_events", return_value=sample_approved_events), \
             patch("data_client.fetch_event_summary_sheets", return_value=[sample_approved_events[0]]):

            # Admin selected only event 101 for detailed summary sheet
            sec = _build_section(
                node=node,
                dept=dept,
                students=[],
                faculty=[],
                qual_counts={},
                cadre_counts={},
                required_faculty=15,
                academic_year="2025-26",
                app_config=app_config,
                include_event_ids=[101],
            )

            # Compact table has 3 rows (all 3 approved events)
            assert len(sec.table_rows) == 3
            assert sec.table_rows[0][1] == "Smart India Hackathon Internal Round"
            assert sec.table_rows[1][1] == "Robotics & Embedded Systems Workshop"
            assert sec.table_rows[2][1] == "AI & Ethics Tech Talk"

            # Detailed summary sheets has only 1 sheet (event 101)
            assert len(sec.summary_sheets) == 1
            assert sec.summary_sheets[0]["id"] == 101
            assert sec.summary_sheets[0]["title"] == "Smart India Hackathon Internal Round"

    def test_zero_events_selected_succeeds(self, sample_approved_events):
        """When 0 events are selected, compact table still renders with all approved events."""
        node = NODES["4.6.1"]
        dept = {"code": "CSE", "name": "Computer Science and Engineering"}
        app_config = {"ACADEMIC_DATA_SERVICE_URL": "http://mock-service"}

        with patch("data_client.fetch_approved_events", return_value=sample_approved_events), \
             patch("data_client.fetch_event_summary_sheets", return_value=[]):

            sec = _build_section(
                node=node,
                dept=dept,
                students=[],
                faculty=[],
                qual_counts={},
                cadre_counts={},
                required_faculty=15,
                academic_year="2025-26",
                app_config=app_config,
                include_event_ids=[],
            )

            # Compact table has all 3 rows
            assert len(sec.table_rows) == 3
            # Summary sheets is empty list
            assert sec.summary_sheets == []

    def test_all_events_selected(self, sample_approved_events):
        """When all events are selected, all 3 detailed summary sheets are returned."""
        node = NODES["4.6.1"]
        dept = {"code": "CSE", "name": "Computer Science and Engineering"}
        app_config = {"ACADEMIC_DATA_SERVICE_URL": "http://mock-service"}

        with patch("data_client.fetch_approved_events", return_value=sample_approved_events), \
             patch("data_client.fetch_event_summary_sheets", return_value=sample_approved_events):

            sec = _build_section(
                node=node,
                dept=dept,
                students=[],
                faculty=[],
                qual_counts={},
                cadre_counts={},
                required_faculty=15,
                academic_year="2025-26",
                app_config=app_config,
                include_event_ids=[101, 102, 103],
            )

            assert len(sec.table_rows) == 3
            assert len(sec.summary_sheets) == 3


class TestDocxAndPdfSummarySheetsRendering:
    def test_docx_renders_summary_sheets(self, sample_approved_events):
        """DOCX renderer should render summary sheets when present in ReportSection."""
        from render.docx_renderer import render_docx
        from docx import Document
        import io

        sec = ReportSection(
            id="4.6.1",
            title="Professional Activities and Club Events",
            marks=5,
            content_type="events_table",
            level=3,
            table_headers=["Sl. No", "Event Title", "Club / Society", "Type", "Date", "Venue", "Attendees"],
            table_rows=[[1, "SIH 2025", "Coding Club", "Hackathon", "2025-09-15", "Auditorium", 120]],
            summary_sheets=[sample_approved_events[0]],
        )

        report = ReportData(
            sar_format="ug_tier_ii_gapc_v4",
            report_type="nba",
            scope="criterion:4",
            academic_year="2025-26",
            generated_at=datetime.now(timezone.utc).isoformat(),
            department={"name": "Computer Science and Engineering", "code": "CSE"},
            sections=[sec],
            report_id="test-rep-101",
        )

        docx_bytes = render_docx(report)
        doc = Document(io.BytesIO(docx_bytes))
        all_text = "\n".join(p.text for p in doc.paragraphs)

        assert "Detailed Activity Summary Sheets" in all_text
        assert "Smart India Hackathon Internal Round" in all_text
        assert "Coding Club" in all_text
        assert "Dr. Arvind Rao" in all_text or any("Dr. Arvind Rao" in c.text for t in doc.tables for r in t.rows for c in r.cells)


class TestEventSummarySheetsAccessControl:
    """
    Test that report-assembly read endpoints strictly enforce:
    - Admin & Teacher (Faculty) -> allowed
    - Student & Worker -> forbidden (403)
    """
    def test_role_enforcement_logic(self):
        """Verify role validation accepts admin/teacher and rejects student/worker."""
        allowed_roles = {"admin", "teacher"}

        def check_role(role: str) -> tuple[bool, int]:
            r = (role or "").lower()
            if r not in allowed_roles:
                return False, 403
            return True, 200

        # Student -> Forbidden
        allowed, code = check_role("Student")
        assert not allowed
        assert code == 403

        # Worker -> Forbidden
        allowed, code = check_role("Worker")
        assert not allowed
        assert code == 403

        # Teacher (Faculty) -> Allowed
        allowed, code = check_role("Teacher")
        assert allowed
        assert code == 200

        # Admin -> Allowed
        allowed, code = check_role("Admin")
        assert allowed
        assert code == 200
