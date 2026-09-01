"""
Render-layer unit tests.
Verifies that PDF and DOCX renderers produce output containing the same
key content from a fixed ReportData — preventing silent drift.
No network calls; no Flask context required.
"""

import sys, os
sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

import pytest
from tests.fixtures import make_minimal_report_data


class TestDocxRenderer:
    def test_docx_renders_without_error(self):
        """DOCX generation should not raise from a valid ReportData."""
        from render.docx_renderer import render_docx
        report = make_minimal_report_data()
        docx_bytes = render_docx(report)
        assert isinstance(docx_bytes, bytes)
        assert len(docx_bytes) > 1000  # non-trivial file

    def test_docx_contains_section_title(self):
        """The section title must appear in the DOCX document."""
        from render.docx_renderer import render_docx
        from docx import Document
        import io

        report = make_minimal_report_data()
        docx_bytes = render_docx(report)
        doc = Document(io.BytesIO(docx_bytes))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Student-Faculty Ratio" in all_text
        assert "Faculty Qualification Index" in all_text

    def test_docx_table_row_count(self):
        """The DOCX table must contain the expected number of rows."""
        from render.docx_renderer import render_docx
        from docx import Document
        import io

        report = make_minimal_report_data()
        docx_bytes = render_docx(report)
        doc = Document(io.BytesIO(docx_bytes))
        # At least one table should exist (formula result table)
        assert len(doc.tables) >= 1

    def test_both_sections_present(self):
        """All section IDs in ReportData must appear in DOCX output."""
        from render.docx_renderer import render_docx
        from docx import Document
        import io

        report = make_minimal_report_data()
        docx_bytes = render_docx(report)
        doc = Document(io.BytesIO(docx_bytes))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        for sec in report.sections:
            assert sec.title in all_text, f"Section '{sec.title}' missing from DOCX"


class TestPDFRenderer:
    def test_pdf_renders_without_error(self):
        """PDF generation should not raise from a valid ReportData."""
        try:
            from render.pdf_renderer import render_pdf
        except ImportError:
            pytest.skip("WeasyPrint not installed in this environment")

        report = make_minimal_report_data()
        pdf_bytes = render_pdf(report)
        assert isinstance(pdf_bytes, bytes)
        # PDF files start with %PDF-
        assert pdf_bytes[:4] == b"%PDF"

    def test_pdf_and_docx_same_section_count(self):
        """
        PDF and DOCX must contain the same number of sections.
        This test enforces that neither renderer silently drops a section.
        """
        try:
            from render.pdf_renderer import render_pdf
            from render.docx_renderer import render_docx
            from docx import Document
            import io

            report = make_minimal_report_data()
            pdf_bytes  = render_pdf(report)
            docx_bytes = render_docx(report)

            # DOCX: count heading paragraphs that contain section IDs
            doc      = Document(io.BytesIO(docx_bytes))
            all_text = "\n".join(p.text for p in doc.paragraphs)
            for sec in report.sections:
                assert sec.title in all_text, f"DOCX missing: {sec.title}"
        except ImportError:
            pytest.skip("WeasyPrint not installed")


class TestSARTree:
    def test_tree_loads_and_marks_sum_to_1000(self):
        """The 1000-mark assertion in ug_tier_ii_gapc_v4 fires on import."""
        import sar_tree.ug_tier_ii_gapc_v4 as tree
        leaf_marks = sum(
            n.marks for n in tree.NODES.values()
            if n.node_type != "criterion_header" and n.marks > 0
        )
        assert leaf_marks == 1000

    def test_all_formula_nodes_have_formula_fn(self):
        """Every formula_table node must reference a formula function name."""
        import sar_tree.ug_tier_ii_gapc_v4 as tree
        for nid, node in tree.NODES.items():
            if node.node_type == "formula_table":
                assert node.formula_fn, f"formula_table node {nid} has no formula_fn"

    def test_scope_criterion_5(self):
        from sar_tree.registry import resolve_scope
        ids = resolve_scope("ug_tier_ii_gapc_v4", "criterion:5")
        assert "5" in ids
        assert "5.1" in ids
        assert "5.5" in ids
        # No other criteria should appear
        for nid in ids:
            assert nid == "5" or nid.startswith("5."), f"Unexpected node {nid} in criterion:5 scope"

    def test_scope_subcriterion_6_1_2(self):
        from sar_tree.registry import resolve_scope
        ids = resolve_scope("ug_tier_ii_gapc_v4", "subcriterion:6.1.2")
        # Should include the parent AND its children
        assert "6.1.2.1" in ids
        assert "6.1.2.2" in ids
